
#Загрузка библиотек для работы с базой данных
import re
import secrets
import shutil
import string

import docx
from sqlalchemy.orm import sessionmaker, relationship
from sqlalchemy import create_engine, ForeignKey
import datetime #Загрузка библиотеки для работы с датой и временем
import os #Загрузка библиотеки для работы с файлами
#Загрузка библиотеки Flask и расширений этой бибилотеки
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, login_required, logout_user, current_user, UserMixin
from flask import Flask, render_template, request, flash, redirect, url_for, send_from_directory, send_file
#Загрузка библиотек для защиты файлов
from werkzeug.security import generate_password_hash, check_password_hash, safe_join

UPLOAD_FOLDER = 'static\\files'
db_string = "postgresql://postgres:123@localhost/CinemaSite" #адрес подключения к БД
dbEngine = create_engine(db_string) #создание класса базы данных

app = Flask( __name__ )
app.secret_key = 'veryveryveryHardSecretKey' #создание ключа для шифрования данных
app.config['SQLALCHEMY_DATABASE_URI'] = db_string #опредления встроенного в Flask базы данных
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER # определение папки для загружаемых файлов


db = SQLAlchemy(app) #Создание экземпляра базы данных
#manager = LoginManager(app) #Создание менеджера логинов
Session = sessionmaker(db) #Создание сессии сайта
session = Session()

db.create_all() #создание элементов базы даных
db.session.commit()

#Функции для работы с базой данных
def get_filmList():
    request = dbEngine.execute(f'SELECT * FROM "filmWithSessionName"') #запрос к базе данных для получения списка фильмов
    filmList = request.fetchall()
    return filmList

def filter_filmList_Genre(genreID): #функция для фильтрации списка фильмов по жанру
    request = dbEngine.execute(f'SELECT * FROM "filmWithSessionName" WHERE genre_id = {genreID}') #запрос к базе данных для вывода фильмов определенного жанра
    filmList = request.fetchall()
    return filmList

def get_filmInfo(filmID): #функция для информации о фильме
    request = dbEngine.execute(f'SELECT * FROM "filmWithSessionName" WHERE id = {filmID}') #запрос к базе данных для получения данных фильма
    filmData = request.fetchone()
    return filmData

def get_filmSessions(filmID): #функция для получения сеансов фильма
    request = dbEngine.execute(f'SELECT * FROM "Sessions" WHERE film_id = {filmID}')  # запрос к базе данных на получение сеансов фильма
    filmSessions = request.fetchall()
    return filmSessions

def get_placesForSession(sesionID):
    session = dbEngine.execute(f'SELECT * FROM "Sessions" WHERE id = {sesionID}').fetchone() #получение данных сеанса в зале
    rowsList = dbEngine.execute(f'SELECT * FROM "Rows" WHERE hall_id = {session.hall_id}').fetchall() #получение рядов в зале
    places = []
    for row in rowsList:
        placesInRow = dbEngine.execute(f'SELECT * FROM "Places" WHERE row_id = {row.id} ORDER BY number').fetchall() #полечение мест в ряду
        # for place in placesInRow:
        places.append(placesInRow) #добавление места в список мест
    return places

#блок функций для обработки адресных путей
@app.route('/', methods=['GET', 'POST'])
@app.route('/home', methods=['GET', 'POST'])
def mainPage():
    return render_template("index.html")

@app.route('/about') #вывод страницы cо справочной информацией
def aboutPage():
    return render_template("about.html")

@app.route('/filmList', methods=['GET', 'POST']) #вывод страницы c афишей кинотеатра
def filmListPage():
    filmList = get_filmList()
    genres = dbEngine.execute('SELECT * FROM "Genres"').fetchall()
    genre_id = request.form.get('genre_id')
    if request.method == "POST" and genre_id:
        print(request.form)
        if   not (genre_id == "all"):
            filmList = filter_filmList_Genre(genre_id)
        # return redirect(url_for('filmListPage',filmList=filmListFiltered,genres=genres))
    return render_template("filmList.html",filmList=filmList,genres=genres)

@app.route('/filmInfo/<filmID>', methods=['GET', 'POST']) #вывод страницы c информацией о выбранном фильме
def filmInfo(filmID):
    filmData = get_filmInfo(filmID)
    print(filmData)
    return render_template("filmInfo.html",filmData=filmData)

@app.route('/sessions/<filmID>', methods=['GET', 'POST']) #вывод страницы с сеансами выбранного фильма
def filmSessions(filmID):
    sessions = get_filmSessions(filmID)
    if request.method == "POST":
        pass
    return render_template("filmSessions.html",sessions=sessions)

@app.route('/placesForSession/<sessionID>', methods=['GET', 'POST']) #вывод страницы c местами в кинозале
def sessionPlaces(sessionID):
    places = get_placesForSession(sessionID)
    # print(places)
    # for placesRow in places:
    #     print(placesRow)

    return render_template("places.html",places=places)

@app.route('/buyTicket/<placeID>', methods=['GET', 'POST']) #вывод страницы покупки билета
def buyTicket(placeID):
    #получение данных из формы
    fio = request.form.get('fio')
    telephone = request.form.get('telephone')

    #получение данных о билете
    place = dbEngine.execute(f'SELECT * FROM "Places" WHERE id = {placeID}').fetchone()
    row = dbEngine.execute(f'SELECT * FROM "Rows" WHERE id = {place.row_id}').fetchone()
    session = dbEngine.execute(f'SELECT * FROM "Sessions" WHERE hall_id = {row.hall_id}').fetchone()
    film = dbEngine.execute(f'SELECT * FROM "Films" WHERE id = {session.film_id}').fetchone()

    if request.method == "POST" and fio and telephone:
        print('НАЧАЛО ПОКУПКИ БИЛЕТА')
        letters_and_digits = string.ascii_letters + string.digits
        password = ''.join(secrets.choice(letters_and_digits) for i in range(8)) #формирование пароля
        code = ''.join(secrets.choice(string.digits) for i in range(6)) #формирование кода для билета
        if dbEngine.execute(f'SELECT * FROM "Buyers" WHERE fio =' +f"'{fio}'").fetchone():
            print('ПОЛЬЗОВАТЕЛЬ ЗАРЕГИСТРИРОВАН')
            buyer = dbEngine.execute(
                f'SELECT * FROM "Buyers" WHERE fio =' + f"'{fio}'").fetchone()  # определение пользователя
        else:
            print('ДОБАВЛЕНИЕ ПОЛЬЗОВАТЕЛЯ')
            dbEngine.execute('INSERT INTO "Buyers" (telephone,password,registered,fio) VALUES ' + f"({telephone},'{password}',{False},'{fio}') RETURNING id") #добавление пользователя
            buyer = dbEngine.execute(f'SELECT * FROM "Buyers" WHERE fio =' +f"'{fio}'").fetchone() #определение пользователя

        print('ДОБАВЛЕНИЕ БИЛЕТА')
        date = datetime.date.today()
        print(date)
        dbEngine.execute('INSERT INTO "Tickets"(code,buyer_id,place_id,session_id,purchase_date) VALUES' + f"({code},{buyer.id},{place.id},{session.id},'{date}') RETURNING id")

        ticket = dbEngine.execute(f'SELECT * FROM "Tickets" WHERE buyer_id = {buyer.id}').fetchone()
        dbEngine.execute(f'UPDATE "Places" SET status = False Where id = {place.id}')
        doc = docx.Document()
        # добавляем первый параграф
        # добавляем еще два параграфа
        par8 = doc.add_paragraph(f'БИЛЕТ {ticket.id}')
        par1 = doc.add_paragraph(f'Ваш код - {ticket.code}')
        par2 = doc.add_paragraph(f'Дата покупки - {ticket.purchase_date}')
        par3 = doc.add_paragraph(f'Место - {place.number}')
        par4 = doc.add_paragraph(f'Дата - {session.date}')
        par5 = doc.add_paragraph(f'Цена - {session.price} РУБЛЕЙ')
        par6 = doc.add_paragraph(f'ФИО - {buyer.fio}')
        par7 = doc.add_paragraph(f'Не забудьте что оплата производится в кинотеатре!')
        par7 = doc.add_paragraph(f'В СЛУЧАЕ УТЕРИ БИЛЕТА ВАМ БУДЕТ НЕОБХОДИМО ВЗЯТЬ С СОБОЙ ПАСПОРТ И ТЕЛЕФОН КОТОРЫЙ ВЫ УКАЗАЛИ ПРИ ПОКУПКЕ БИЛЕТА')

        # добавляем текст во второй параграф
        # добавляем текст в третий параграф
        print(os.getcwd())
        os.chdir("static")
        os.chdir("files")
        print(os.getcwd())

        doc.save(f'ticket{ticket.id}.docx')
        os.chdir(os.pardir)
        os.chdir(os.pardir)
        print(os.getcwd())
        # shutil.move("/ticket{ticket.id}.docx", app.config["UPLOAD_FOLDER"])
        filename = f"ticket{ticket.id}.docx"

        return redirect(url_for('download_file',name=filename))

    return render_template("buy.html",place=place,film=film, row=row, session=session)

@app.route('/uploads/<name>') #загрузка файла на сервер
def download_file(name):
    # print(send_from_directory(app.config["UPLOAD_FOLDER"], name)) #вывод информации о файле в консоль
    return send_file(safe_join('static/files', name),as_attachment=True)


if __name__ =="__main__":
    app.run(host=os.getenv('IP', '0.0.0.0'),
            port=int(os.getenv('PORT', 4444)),debug=True)



